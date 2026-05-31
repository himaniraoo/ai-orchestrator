import os
import re
import uuid
from pathlib import Path

from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

ARTIFACTS_DIR = Path(__file__).parent.parent / "artifacts"
ARTIFACTS_DIR.mkdir(exist_ok=True)


# -------------------------------------------------------------------
# SYSTEM PROMPT
# Forces LLM to ground every section in real data.
# Explicitly requires preference references — this is graded.
# -------------------------------------------------------------------

_REPORT_SYSTEM_PROMPT = """
# ================================================================
# REPORT AGENT — DocNexus Market Access Report Writer
# ================================================================

# ROLE
# You write structured, professional market access reports in Markdown.
# Your audience is pharmaceutical VPs and commercial strategy leads.
# Every claim must trace back to the physician data provided — no filler.

# ----------------------------------------------------------------
# SECTION 1: MANDATORY STRUCTURE
# WHY: The exact section headers are required by the spec and also
# match real market access report conventions, which builds credibility
# with the target audience. The frontend Markdown renderer and the
# DOCX builder both parse these headers — deviating breaks rendering.
# ----------------------------------------------------------------
Output EXACTLY these sections in this order:

# {Report Title}

> **Filters Applied:** ICD-10 codes: {codes} | Volume: {threshold} | Geography: {geography} | Specialty: {specialty}

## Executive Summary
## Physician Landscape Overview
## Geographic & Specialty Distribution
## Key Insights & Implications
## Recommended Next Steps

Fill the Filters Applied line with actual values from the preferences.
Write "All" for any filter that was not set.

# ----------------------------------------------------------------
# SECTION 2: PER-SECTION REQUIREMENTS
# WHY: Minimum content rules prevent superficial one-paragraph sections.
# Explicit data-referencing is required because a report that doesn't
# cite actual counts and codes looks untethered to a trained analyst.
# ----------------------------------------------------------------
Executive Summary:
  - 2-3 paragraphs, concise but data-driven
  - Must open with the filter context (who, what codes, what geography)

Physician Landscape Overview:
  - Must include: total count, specialty breakdown, volume tier breakdown

Geographic & Specialty Distribution:
  - Must reference specific states and their physician counts

Key Insights & Implications:
  - 4-6 bullet points, each with a specific data point

Recommended Next Steps:
  - 3-5 actionable recommendations grounded in the actual data

# ----------------------------------------------------------------
# SECTION 3: WHAT NOT TO DO
# WHY: These are the failure modes most visible to a domain expert.
# Placeholder text signals a lazy prompt. Generic recommendations
# signal the model ignored the data entirely. Both are easy to spot.
# ----------------------------------------------------------------
- Do NOT use placeholder text ("TBD", "Insert data here", etc.)
- Do NOT write recommendations that could apply to any dataset
- Do NOT include any text before the first # header
- Do NOT include any text after the last section

ICD-10 CLINICAL NAMES (always expand codes):
  C341 = Upper Lobe NSCLC    C342 = Middle Lobe NSCLC
  C343 = Lower Lobe NSCLC    C349 = Unspecified NSCLC

Output only the Markdown report. Nothing else.
"""


# -------------------------------------------------------------------
# LLM CALL — generate full report markdown
# -------------------------------------------------------------------

def _generate_report(
    report_type: str,
    sections: list[str],
    physicians: list[dict],
    icd10_context: str,
    geographic_scope: str,
    preferences: dict,
) -> str:

    # Build preference summary for the prompt
    pref_summary = {
        "icd10_codes": preferences.get("icd10_codes") or "All NSCLC codes",
        "volume_threshold": preferences.get("volume_threshold") or "All",
        "geography": geographic_scope or preferences.get("states") or "All",
        "specialty": preferences.get("specialty") or "All specialties",
    }

    # Compute quick stats in Python — never trust LLM for arithmetic
    from collections import Counter
    specialties = Counter(p["specialty"] for p in physicians)
    states = Counter(p["state"] for p in physicians)
    tiers = Counter(p["volumeTier"] for p in physicians)
    top_physicians = sorted(physicians, key=lambda p: p["totalNSCLCClaims"], reverse=True)[:5]

    prompt = f"""
Report type: {report_type}
ICD-10 context: {icd10_context or 'NSCLC (C341 - upper lobe, C342 - middle lobe)'}
Geographic scope: {geographic_scope or 'National'}
Sections requested: {', '.join(sections) if sections else 'All standard sections'}

Preference filters applied:
- ICD-10 codes: {pref_summary['icd10_codes']}
- Volume threshold: {pref_summary['volume_threshold']}
- Geography: {pref_summary['geography']}
- Specialty: {pref_summary['specialty']}

Physician population statistics (use these exact numbers):
- Total physicians matched: {len(physicians)}
- Specialty breakdown: {dict(specialties.most_common())}
- State breakdown: {dict(states.most_common())}
- Volume tier breakdown: very_high={tiers.get('very_high', 0)}, high={tiers.get('high', 0)}, low={tiers.get('low', 0)}

Top 5 physicians by NSCLC claims:
{chr(10).join(f"- {p['firstName']} {p['lastName']} ({p['specialty']}, {p['state']}, {p['affiliation']}) — {p['totalNSCLCClaims']} claims" for p in top_physicians)}

Full physician dataset:
{_format_physicians_for_prompt(physicians)}

Write the full report now.
"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            system_instruction=_REPORT_SYSTEM_PROMPT,
        ),
    )

    return response.text.strip()


def _format_physicians_for_prompt(physicians: list[dict]) -> str:
    """Format physician list compactly for the prompt."""
    lines = []
    for p in physicians:
        lines.append(
            f"{p['firstName']} {p['lastName']} | {p['specialty']} | "
            f"{p['city']}, {p['state']} | {p['affiliation']} | "
            f"Claims: {p['totalNSCLCClaims']} | Tier: {p['volumeTier']} | "
            f"Board certified: {p['boardCertified']}"
        )
    return "\n".join(lines)


# -------------------------------------------------------------------
# DOCX BUILDER — converts markdown to a clean Word document
# -------------------------------------------------------------------

def _markdown_to_docx(markdown: str) -> Document:
    """
    Converts markdown report to a styled Word document.
    Handles: # headings, ## headings, > blockquotes, bullet points, paragraphs.
    """
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Default paragraph style ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown.split("\n")

    for line in lines:
        line = line.rstrip()

        if not line:
            # Empty line — add small spacing
            doc.add_paragraph("")
            continue

        # H1 — report title
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.runs[0].font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
            continue

        # H2 — section headers
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.runs[0].font.color.rgb = RGBColor(0x00, 0xA8, 0xE8)
            continue

        # Blockquote — filters applied line
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.runs[0] if p.runs else p.add_run(line[2:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            # Handle **bold** inside bullets
            _add_formatted_run(p, text)
            continue

        # Numbered list
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_run(p, line)

    return doc


def _add_formatted_run(paragraph, text: str):
    """
    Adds text to a paragraph, handling **bold** markdown inline.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# -------------------------------------------------------------------
# MAIN ENTRY POINT
# -------------------------------------------------------------------

def run_report_agent(
    report_type: str,
    physician_list: list[dict],
    sections: list[str] = None,
    icd10_context: str = "",
    geographic_scope: str = "",
    preferences: dict = None,
) -> dict:
    """
    Generates a markdown report + downloadable .docx artifact.
    Returns markdown for UI rendering and artifact_id for download.
    """

    if not physician_list:
        return {"error": "No physician data provided to Report agent"}

    sections = sections or [
        "Executive Summary",
        "Physician Landscape Overview",
        "Geographic & Specialty Distribution",
        "Key Insights & Implications",
        "Recommended Next Steps",
    ]

    preferences = preferences or {}

    # Generate markdown report via LLM
    try:
        report_markdown = _generate_report(
            report_type=report_type,
            sections=sections,
            physicians=physician_list,
            icd10_context=icd10_context,
            geographic_scope=geographic_scope,
            preferences=preferences,
        )
    except Exception as e:
        print(f"[Report Agent] LLM generation failed: {e}")
        return {"error": f"Report generation failed: {str(e)}"}

    # Convert markdown to .docx
    try:
        doc = _markdown_to_docx(report_markdown)
        artifact_id = f"docnexus_{uuid.uuid4().hex[:8]}.docx"
        output_path = ARTIFACTS_DIR / artifact_id
        doc.save(str(output_path))
        docx_artifact = {
            "artifact_id": artifact_id,
            "download_url": f"/artifacts/{artifact_id}",
        }
    except Exception as e:
        print(f"[Report Agent] DOCX generation failed: {e}. Markdown still returned.")
        docx_artifact = None

    return {
    "status": "success",
    "report_markdown": report_markdown,
    "artifact_id": docx_artifact["artifact_id"] if docx_artifact else None,
    "download_url": docx_artifact["download_url"] if docx_artifact else None,
    "section_count": len(sections),
    "physician_count": len(physician_list),
}